import streamlit as st
import os, uuid, fitz, json, time, re
import pandas as pd
from io import BytesIO
from sqlalchemy import create_engine, Column, String, Text, func
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor

# --- CONFIGURAZIONE ---
AUTORE_SOFTWARE = "Sebastiano Barusco"
COPYRIGHT_NOTE = "© 2026 Sebastiano Barusco - Tutti i diritti riservati"
DB_URL = "sqlite:///./osservatorio.db"

def trova_logo():
    for file in os.listdir('.'):
        if file.lower().startswith('logo') and any(file.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg']):
            return file
    return None

LOGO_PATH = trova_logo()

st.set_page_config(page_title="Osservatorio ACDT", page_icon="⚖️", layout="wide")

# Style istituzionale Bordeaux ACDT
st.markdown("""
    <style>
    .main { background-color: #fcfcfc; }
    h1, h2, h3, h4 { color: #8a1c3d !important; }
    .stButton>button { background-color: #8a1c3d !important; color: white !important; border-radius: 8px; font-weight: bold; }
    .stMetric { background-color: #ffffff; padding: 15px; border-radius: 10px; border-left: 5px solid #8a1c3d; box-shadow: 0 2px 4px rgba(0,0,0,0.05); }
    .footer { text-align: center; color: #6c757d; padding: 20px; border-top: 1px solid #dee2e6; margin-top: 50px; font-size: 0.8rem; }
    </style>
    """, unsafe_allow_html=True)

# --- DATABASE ---
engine = create_engine(DB_URL, connect_args={"check_same_thread": False})
Base = declarative_base()
class Sentenza(Base):
    __tablename__ = "sentenze"
    id = Column(String, primary_key=True)
    stato = Column(String, default="Nuovo") 
    organo = Column(String, default="N/D")
    numero = Column(String, default="N/D")
    massima = Column(Text, default="N/D")
    autore = Column(String, default="Redazione")
    file_path = Column(String)

Base.metadata.create_all(bind=engine)
SessionLocal = sessionmaker(bind=engine)
db = SessionLocal()

# --- UTILS AI ROBUSTE ---
def formatta_massima_sistematica(m_dati):
    """Estrae i valori dal dizionario con OGGETTO AMPLIATO e senza Keywords"""
    def get_val(keys):
        for k in keys:
            if k in m_dati: return m_dati[k]
            if k.lower() in m_dati: return m_dati[k.lower()]
            if k.replace(" ", "_") in m_dati: return m_dati[k.replace(" ", "_")]
        return "Dato non rilevato"

    testo = f"**OGGETTO DELLA CAUSA (FATTISPECIE)**:\n{get_val(['oggetto_ampliato', 'oggetto', 'descrizione_causa'])}\n\n"
    testo += f"**PRINCIPIO DI DIRITTO**: \n{get_val(['principio', 'principio di diritto'])}\n\n"
    testo += f"**RIFERIMENTI NORMATIVI**: \n{get_val(['norme', 'riferimenti normativi'])}\n\n"
    testo += f"**ESITO DELLA DECISIONE**: \n{get_val(['esito', 'decisione'])}"
    return testo

def pulisci_e_carica_json(testo_raw):
    try:
        testo_pulito = re.sub(r'```json\s*|```', '', testo_raw).strip()
        match = re.search(r'\{.*\}', testo_pulito, re.DOTALL)
        return json.loads(match.group()) if match else json.loads(testo_pulito)
    except: return None

def analizza_sentenza(file_path):
    api_key = os.getenv("GROQ_API_KEY")
    if not api_key: return None, "Manca la chiave API."
    client = Groq(api_key=api_key)
    try:
        doc = fitz.open(file_path)
        # Leggiamo un contesto più ampio per l'oggetto dettagliato (prime 5 pagine + ultima)
        testo_estratto = "\n".join([doc[i].get_text() for i in range(min(5, len(doc)))]) + "\n" + doc[-1].get_text()
        doc.close()
        
        prompt = f"""Sei un Magistrato addetto all'Ufficio del Massimario. Analizza la sentenza ed estrai i dati in JSON puro.
        CAMPI RICHIESTI:
        1. "organo": Nome corte.
        2. "numero": Numero sentenza/anno.
        3. "massima_dati": {{ 
            "oggetto_ampliato": "Descrizione analitica e dettagliata della lite, dei motivi del ricorso e dei fatti di causa.", 
            "principio": "Enunciazione del principio di diritto applicato (ratio decidendi).", 
            "norme": "Articoli di legge rilevanti citati.", 
            "esito": "Sintesi dell'esito della decisione." 
        }}
        REGOLE: Sii molto dettagliato nell'oggetto della causa. Rispondi esclusivamente in JSON puro.
        TESTO: {testo_estratto[:9000]}"""

        chat = client.chat.completions.create(
            messages=[{"role": "system", "content": "Sei un redattore esperto di massimari giuridici. Rispondi in JSON."},
                      {"role": "user", "content": prompt}],
            model="llama-3.1-8b-instant", temperature=0.1
        )
        dati = pulisci_e_carica_json(chat.choices[0].message.content)
        if not dati: return None, "Errore formato JSON."
        
        return {
            "organo": dati.get("organo", "N/D"),
            "numero": dati.get("numero", "N/D"),
            "massima": formatta_massima_sistematica(dati.get("massima_dati", dati))
        }, None
    except Exception as e: return None, str(e)

def ricerca_ai(domanda, sentenze_validate):
    api_key = os.getenv("GROQ_API_KEY")
    client = Groq(api_key=api_key)
    context = "\n".join([f"ID:{s.id} | INFO:{s.organo} n.{s.numero} | TESTO:{s.massima[:700]}" for s in sentenze_validate])
    prompt = f"""L'utente cerca: "{domanda}". Identifica tra queste sentenze le più rilevanti per risolvere il quesito. 
    Restituisci JSON: {{"risultati": [{{"id": "...", "perche": "..."}}]}}
    LISTA PRECEDENTI: {context}"""
    try:
        chat = client.chat.completions.create(
            messages=[{"role": "system", "content": "Sei un assistente legale specializzato in tributario. Rispondi in JSON."},
                      {"role": "user", "content": prompt}],
            model="llama-3.1-8b-instant", temperature=0.1
        )
        return pulisci_e_carica_json(chat.choices[0].message.content)
    except: return None

# --- ESPORTAZIONI ---
def genera_word(lista_sentenze):
    doc = Document()
    doc.add_heading('Massimario ACDT - Report Giurisprudenziale', 0).alignment = 1
    for i in lista_sentenze:
        p = doc.add_paragraph()
        run = p.add_run(f"{i.organo}")
        run.bold, run.font.size, run.font.color.rgb = True, Pt(12), RGBColor(138, 28, 61)
        doc.add_paragraph(f"Sentenza n. {i.numero} | Redattore: {i.autore}").italic = True
        doc.add_paragraph(i.massima.replace("**", "")).alignment = 3
        doc.add_paragraph("-" * 20).alignment = 1
    target = BytesIO(); doc.save(target)
    return target.getvalue()

# --- INTERFACCIA ---
with st.sidebar:
    if LOGO_PATH: st.image(LOGO_PATH, use_container_width=True)
    st.markdown("---")
    st.markdown(f"👨‍💻 **Author:** {AUTORE_SOFTWARE}")
    st.caption(COPYRIGHT_NOTE)

tab_home, tab_gest, tab_search, tab_arch = st.tabs(["🏠 Home Page", "📋 Gestione Analisi", "🔍 Ricerca Intelligente AI", "📚 Archivio Storico"])

# --- TAB HOME ---
with tab_home:
    col_l, col_r = st.columns([0.25, 0.75])
    if LOGO_PATH: col_l.image(LOGO_PATH, use_container_width=True)
    with col_r:
        st.markdown("# Osservatorio Giurisprudenza Tributaria")
        st.markdown("### Associazione Commercialisti Difensori Tributari del Veneto")
    st.markdown("---")
    total = db.query(Sentenza).count()
    validati = db.query(Sentenza).filter(Sentenza.stato == "Validato").count()
    c_s1, c_s2, c_s3 = st.columns(3)
    c_s1.metric("Sentenze Processate", total)
    c_s2.metric("In Archivio Storico", validati)
    c_s3.metric("Ricerca AI", "Attiva")
    st.markdown("---")
    c1, c2 = st.columns(2)
    with c1:
        st.markdown("#### 🔍 Come Funziona?")
        st.write("""
        1. **Caricamento**: Trascina PDF multipli nella sezione 'Gestione'.
        2. **Analisi IA**: Il sistema scansiona il testo e produce una bozza con l'oggetto della causa ampliato e il principio di diritto.
        3. **Validazione**: Conferma o modifica le bozze generate.
        4. **Ricerca AI**: Interroga l'intero archivio utilizzando il linguaggio naturale.
        """)
    with c2:
        st.markdown("#### ⚖️ Metodologia Scientifica")
        st.write("""
        La redazione delle massime segue uno schema analitico:
        - **Oggetto della causa**: Descrizione dettagliata dei fatti e della lite.
        - **Principio di Diritto**: La sintesi della 'ratio decidendi'.
        - **Norme**: Riferimenti legislativi.
        - **Esito**: Verdetto finale.
        """)
    st.info("Piattaforma tecnologica sviluppata per l'Associazione Commercialisti Difensori Tributari del Veneto.")

# --- TAB GESTIONE ---
with tab_gest:
    col_up, col_rev = st.columns([0.35, 0.65])
    with col_up:
        st.subheader("📤 Caricamento PDF")
        firma = st.text_input("Firma Caricamento", value="Redazione")
        u_files = st.file_uploader("Seleziona sentenze", type="pdf", accept_multiple_files=True)
        if st.button("🚀 AVVIA ANALISI DETTAGLIATA"):
            if u_files:
                for f in u_files:
                    f_id = str(uuid.uuid4()); path = f"storage/{f_id}.pdf"
                    os.makedirs("storage", exist_ok=True)
                    with open(path, "wb") as out: out.write(f.getbuffer())
                    with st.spinner(f"Analisi fattispecie: {f.name}..."):
                        res, err = analizza_sentenza(path)
                        if not err:
                            db.add(Sentenza(id=f_id, organo=res["organo"], numero=res["numero"], massima=res["massima"], autore=firma, file_path=path))
                            db.commit()
                st.rerun()
    with col_rev:
        st.subheader("✍️ Revisione Massime")
        nuovi = db.query(Sentenza).filter(Sentenza.stato == "Nuovo").all()
        if nuovi:
            if st.button("🚀 PUBBLICA TUTTI I DOCUMENTI"):
                for s in nuovi: s.stato = "Validato"
                db.commit(); st.rerun()
            for s in nuovi:
                with st.expander(f"📝 {s.organo} - n. {s.numero}", expanded=True):
                    o = st.text_input("Corte", s.organo, key=f"o{s.id}")
                    n = st.text_input("N.", s.numero, key=f"n{s.id}")
                    m = st.text_area("Bozza Massima (Oggetto Ampliato)", s.massima, height=350, key=f"m{s.id}")
                    if st.button("✅ VALIDA E ARCHIVIA", key=f"p{s.id}"):
                        s.organo, s.numero, s.massima, s.stato = o, n, m, "Validato"
                        db.commit(); st.rerun()
        else: st.info("Nessuna analisi da revisionare.")

# --- TAB RICERCA AI ---
with tab_search:
    st.subheader("🔍 Ricerca Semantica con Intelligenza Artificiale")
    domanda = st.text_input("Descrivi un caso o poni un quesito giuridico...")
    if domanda:
        sv = db.query(Sentenza).filter(Sentenza.stato == "Validato").all()
        if sv:
            with st.spinner("L'IA sta consultando i precedenti in archivio..."):
                r_ia = ricerca_ai(domanda, sv)
                if r_ia and r_ia.get("risultati"):
                    for res in r_ia["risultati"]:
                        s = db.query(Sentenza).filter(Sentenza.id == res["id"]).first()
                        if s:
                            with st.container(border=True):
                                st.markdown(f"#### {s.organo} - n. {s.numero}")
                                st.success(f"**RILEVANZA**: {res['perche']}")
                                with st.expander("Leggi Massima"): st.write(s.massima)
                else: st.warning("Nessun precedente trovato per questo quesito.")
        else: st.warning("L'archivio è vuoto.")

# --- TAB ARCHIVIO ---
with tab_arch:
    st.subheader("📚 Archivio Storico")
    arch = db.query(Sentenza).filter(Sentenza.stato == "Validato").all()
    if arch:
        c1, c2, c3 = st.columns([0.4, 0.3, 0.3])
        out_ex = BytesIO()
        with pd.ExcelWriter(out_ex, engine='openpyxl') as wr:
            pd.DataFrame([{"Organo": i.organo, "Numero": i.numero, "Massima": i.massima.replace("**",""), "Autore": i.autore} for i in arch]).to_excel(wr, index=False)
        c1.download_button("📊 EXCEL RIEPILOGO", out_ex.getvalue(), "archivio.xlsx", use_container_width=True)
        c2.download_button("📝 WORD REPORT", genera_word(arch), "archivio.docx", use_container_width=True)
        if c3.button("⚠️ RESET"):
            db.query(Sentenza).delete(); db.commit(); st.rerun()
        st.divider()
        sk = st.text_input("🔍 Filtro rapido (parole chiave)...")
        for i in arch:
            if sk.lower() in i.massima.lower() or sk.lower() in i.organo.lower():
                with st.container(border=True):
                    ca, cb = st.columns([0.85, 0.15])
                    with ca:
                        st.markdown(f"#### {i.organo}\n**n. {i.numero}**")
                        st.write(i.massima)
                    with cb:
                        if os.path.exists(i.file_path):
                            with open(i.file_path, "rb") as fp:
                                st.download_button("📄 PDF", fp, f"sentenza_{i.numero}.pdf", key=f"dl_{i.id}")
    else: st.info("Archivio vuoto.")

st.markdown(f"<div class='footer'>{COPYRIGHT_NOTE} | {AUTORE_SOFTWARE}</div>", unsafe_allow_html=True)
db.close()
